import os
import sys
import zipfile
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt
from abc import abstractmethod


def read_textfile(fpath):
    if not os.path.isfile(fpath):
        raise FileNotFoundError(f"{fpath} is not a file, pwd = {os.getcwd()}")

    with open(fpath, 'r') as f:
        content = f.read()
    return content


def read_docxfile(fpath):
    return "\n".join(p.text for p in Document(fpath).paragraphs)


def read_pdf(fpath):
    from tika import parser
    raw = parser.from_file(fpath)
    return raw['content'].strip()


def safe_mkdir(directory):
    Path(directory).mkdir(parents=True, exist_ok=True)


def extract_zip(src, dest):
    with zipfile.ZipFile(src) as f:
        f.extractall(path=dest)


def zipdir(src, dest):
    relroot = os.path.abspath(os.path.join(src, os.pardir))
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zip:
        for root, dirs, files in os.walk(src):
            # add directory (needed for empty dirs)
            zip.write(root, os.path.relpath(root, relroot))
            for file in files:
                filename = os.path.join(root, file)
                if os.path.isfile(filename):  # regular files only
                    arcname = os.path.join(os.path.relpath(root, relroot), file)
                    zip.write(filename, arcname)


class Writer:
    @abstractmethod
    def write(self, content, fname):
        pass


class DocxWriter(Writer):
    def write(self, content, fname):
        if not fname.endswith(".docx"):
            fname += ".docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        for para in content.split("\n"):
            doc.add_paragraph(para)

        doc.save(fname)


def docx_insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


class DocxInsertionWriter(Writer):
    def __init__(self, template_path, pre_para_id, insert_before=True):
        self.template_path = template_path
        self.pre_para_id = pre_para_id
        self.insert_before = insert_before

    def write(self, content, fname):
        if not fname.endswith(".docx"):
            fname += ".docx"
        doc = Document(self.template_path)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        pre_para = doc.paragraphs[self.pre_para_id]

        if self.insert_before:
            for para_text in content.split("\n"):
                pre_para.insert_paragraph_before(text=para_text, style=style)
        else:
            for para_text in content.split("\n")[::-1]:
                docx_insert_paragraph_after(pre_para, text=para_text, style=style)

        print(f"saving to {fname}")
        doc.save(fname)
